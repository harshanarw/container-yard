"""
Generate: MnR Estimate Calculation Guide (Word Document)
Run: python3 generate_estimate_guide.py
Output: MnR_Estimate_Calculation_Guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.0)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0, 70, 127)):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(*color)
    p.runs[0].font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}[level])
    return p

def para(text='', bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size, color=color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def shade_cell(cell, hex_color='D9E1F2'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, header_fill='1F4E79', header_text='FFFFFF'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*bytes.fromhex(header_text))
        shade_cell(cell, header_fill)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.add_row()
        fill = 'EEF2F7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            shade_cell(cell, fill)

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return t

def box_note(text, fill='FFF2CC', border='FFC000'):
    """Yellow info box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    run = p.add_run(f'  📌  {text}')
    run.font.size = Pt(9.5)
    run.italic = True
    return p

def example_box(title, lines):
    """Grey example block."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'  Example — {title}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 112, 192)

    for line in lines:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.3)
        run2 = p2.add_run(f'      {line}')
        run2.font.name = 'Courier New'
        run2.font.size = Pt(9)
        shade_cell_p(p2, 'F2F2F2')
    doc.add_paragraph()

def shade_cell_p(para, hex_color):
    """Apply shading to a paragraph (via pPr)."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONTAINER YARD MANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 70, 127)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Maintenance & Repair (MnR)')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Estimate Calculation Guide')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 55)
run.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('How Survey Damage Findings are converted into\nLabour Hours, Labour Cost and Material Cost\nfor Customer Approval')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(68, 68, 68)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('For internal use — Operations & Finance teams')
run.font.size = Pt(10)
run.italic = True
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading('1.  What Is an MnR Estimate?', 1)

para(
    'When a shipping container arrives at the yard and is found to have damage, '
    'the yard prepares a Maintenance & Repair (MnR) Estimate. This estimate lists '
    'every repair needed, how many labour hours each job will take, and what materials '
    'are required — along with the total cost. The estimate is then sent to the container '
    'owner (the shipping line or agent) for approval before any repair work begins.',
    size=11
)

doc.add_paragraph()
para('The process has four stages:', bold=True)
numbered('Survey — the yard inspector walks around the container and records each damage found.')
numbered('Tariff lookup — the system finds the matching repair rate from the MnR Tariff master.')
numbered('Estimate creation — the system builds the estimate line items with labour and material costs.')
numbered('Customer approval — the estimate is sent to the customer through the owner portal or by email.')

doc.add_paragraph()
box_note(
    'Think of the MnR Tariff as a "price list" for every type of repair. '
    'The survey tells the system what is damaged; the tariff tells it how much the repair costs.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MnR CODE SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading('2.  Understanding MnR Codes', 1)

para(
    'Every damage and repair is described using a set of standard codes. '
    'These codes make it possible to look up the correct tariff rate automatically.',
    size=11
)

add_table(
    headers=['Code Type', 'What It Describes', 'Examples'],
    rows=[
        ['Location Code',   'Where on the container the damage is',         'DOOR, ROOF, SIDE-L, FLOOR'],
        ['Component Code',  'Which part is damaged',                         'PANEL, POST, RAIL, LOCK, HINGE'],
        ['Damage Code',     'What kind of damage it is',                     'DENT, HOLE, CRACKED, CORRODED'],
        ['Repair Code',     'What repair is needed',                         'STR (Straighten), RPL (Replace), WLD (Weld), TAP (Paint)'],
        ['Material Code',   'What material is used for the repair',          'STEEL-PLATE, PAINT-GREY, TIMBER'],
        ['Responsibility',  'Who is responsible for the repair cost',        'OWNER, LESSEE, YARD'],
    ],
    col_widths=[1.3, 2.5, 2.5],
    header_fill='1F4E79'
)

para(
    'These codes are used in two places: (1) the inspector fills them in on the survey damage record, '
    'and (2) the tariff master uses them to define repair rates.',
    size=11, italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MnR TARIFF STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('3.  MnR Tariff Structure', 1)

para(
    'The MnR Tariff master is organised in three layers. Understanding these layers '
    'is the key to understanding how every estimate amount is calculated.',
    size=11
)

# 3.1 Tariff Header
heading('3.1  Tariff Header (the "Agreement")', 2)

para(
    'The Tariff Header is the top-level record. It represents a rate agreement — '
    'either a customer-specific agreement or a general default rate for all customers.',
    size=11
)

add_table(
    headers=['Field', 'Meaning'],
    rows=[
        ['Name',                'Friendly name, e.g. "Standard Tariff 2025" or "Evergreen Line Tariff"'],
        ['Customer',            'If blank, this is the global default. If set, it overrides the default for that customer.'],
        ['Valid From / Valid To','The date range during which this tariff applies.'],
        ['Currency',            'The base currency for all rates in this tariff (usually USD).'],
        ['Applicable Sizes',    'Which container sizes this tariff covers: 20ft, 40ft, 45ft, etc.'],
        ['Active',              'Only active tariffs are used for calculations.'],
    ],
    col_widths=[1.8, 4.5],
    header_fill='2E75B6'
)

box_note(
    'Rule: If a customer-specific tariff exists for a container owner, it is used. '
    'Otherwise the system falls back to the global default tariff.'
)

# 3.2 Tariff Rules
heading('3.2  Tariff Rules (Direct Rate Lookup)', 2)

para(
    'A Tariff Rule defines the exact labour and material rate for a specific '
    'combination of Component Code + Damage Code + Repair Code. '
    'This is the most common and simplest method for setting repair prices.',
    size=11
)

add_table(
    headers=['Field', 'Meaning', 'Example'],
    rows=[
        ['Component Code',    'Which part',                     'PANEL'],
        ['Damage Code',       'Type of damage (optional)',       'DENT'],
        ['Repair Code',       'Type of repair',                  'STR (Straighten)'],
        ['Std Labour Hours',  'Standard hours for this repair',  '1.50 hrs'],
        ['Labour Rate',       'Cost per labour hour (USD)',       'USD 25.00/hr'],
        ['Material Qty',      'Quantity of material needed',      '0 (no material for straightening)'],
        ['Material Rate',     'Cost per unit of material (USD)',  'USD 0.00'],
        ['Ancillary',         'Any extra fixed charge',           'USD 5.00'],
        ['Min Charge',        'Minimum amount even if calc is lower', 'USD 20.00'],
        ['Max Charge',        'Maximum amount even if calc is higher','USD 200.00'],
    ],
    col_widths=[1.6, 2.4, 1.8],
    header_fill='2E75B6'
)

example_box(
    'Straighten a Side Panel (PANEL + STR)',
    [
        'Std Labour Hours : 1.50',
        'Labour Rate      : USD 25.00 / hr',
        'Material Qty     : 0',
        'Material Rate    : USD 0.00',
        'Ancillary        : USD 0.00',
        '',
        'Labour Amount    = 1.50 × 25.00 = USD 37.50',
        'Material Amount  = 0 × 0.00    = USD 0.00',
        'Total            = USD 37.50',
    ]
)

example_box(
    'Replace a Floor Board (FLOOR + RPL)',
    [
        'Std Labour Hours : 2.00',
        'Labour Rate      : USD 25.00 / hr',
        'Material Qty     : 3  (boards)',
        'Material Rate    : USD 18.00 / board',
        'Ancillary        : USD 0.00',
        '',
        'Labour Amount    = 2.00 × 25.00 = USD 50.00',
        'Material Amount  = 3 × 18.00   = USD 54.00',
        'Total            = USD 104.00',
    ]
)

# 3.3 Tariff Items & Slabs
heading('3.3  Tariff Items and Slabs (Quantity-Based Pricing)', 2)

para(
    'Some repairs are priced differently depending on how many units need to be repaired. '
    'For example, replacing 1 door lock is expensive per unit, but if 5 locks need replacing, '
    'the rate per lock might drop. This is handled by Tariff Items and Slabs.',
    size=11
)

para('A Tariff Item defines an operation (e.g., "Replace Twist Lock"). Each item has multiple Slabs:', bold=True, size=11)

add_table(
    headers=['Slab Field', 'Meaning'],
    rows=[
        ['Slab Label',     'A friendly name for this tier, e.g. "1–5 units" or "6–10 units"'],
        ['Qty From',       'The minimum quantity that triggers this slab. E.g. qty_from=1 means "for 1 or more items"'],
        ['Labour Hours',   'Labour hours for this tier\'s base quantity'],
        ['Material Cost',  'Material cost for this tier\'s base quantity'],
        ['Is Additional',  'If ticked, this row is the "per extra unit" rate beyond the base slab'],
    ],
    col_widths=[1.6, 4.7],
    header_fill='2E75B6'
)

para('How the system picks the right slab:', bold=True)
numbered('It looks at all slabs where Qty From ≤ the actual damaged quantity.')
numbered('It picks the slab with the highest Qty From value (the "best match").')
numbered('If there is an "Additional" slab, extra units beyond the base slab are priced using it.')

example_box(
    'Replace Twist Locks — Slab Pricing (qty = 8 locks)',
    [
        'Slab Setup:',
        '  Slab 1  — qty_from: 1,  labour_hours: 0.25,  material_cost: USD 12.00  (is_additional: No)',
        '  Slab 2  — qty_from: 5,  labour_hours: 1.00,  material_cost: USD 50.00  (is_additional: No)',
        '  Extra   — qty_from: 1,  labour_hours: 0.20,  material_cost: USD 10.00  (is_additional: Yes)',
        '',
        'Actual Quantity = 8 locks',
        '',
        'Step 1 — Find best base slab:',
        '  qty=8 >= qty_from=1 ✓   (Slab 1 matches)',
        '  qty=8 >= qty_from=5 ✓   (Slab 2 matches) ← higher Qty From wins',
        '  Base Slab = Slab 2  →  hours=1.00,  material=USD 50.00',
        '',
        'Step 2 — Calculate extras:',
        '  Extra units = 8 - 5 = 3 units',
        '  Multiplier  = 3 / 1  = 3',
        '  Extra hours   = 3 × 0.20  = 0.60',
        '  Extra material = 3 × 10.00 = USD 30.00',
        '',
        'Step 3 — Add base + extras:',
        '  Total Hours    = 1.00 + 0.60 = 1.60 hrs',
        '  Total Material = 50.00 + 30.00 = USD 80.00',
        '',
        'Step 4 — Calculate Labour Cost:',
        '  Labour Rate   = USD 25.00 / hr',
        '  Labour Amount = 1.60 × 25.00 = USD 40.00',
        '',
        'Final Line Total = USD 40.00 + USD 80.00 = USD 120.00',
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — FROM SURVEY TO ESTIMATE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('4.  From Survey Damage to Estimate Line Item', 1)

para(
    'When the inspector completes a survey and the estimate is created, '
    'the system automatically converts each damage record into an estimate line item '
    'by following these steps:',
    size=11
)

numbered(
    'Load all damages from the survey (inquiry record), each with its location, component, '
    'damage type, repair type, material, and quantity.'
)
numbered(
    'Find the best tariff: if the container owner has a customer-specific tariff, use it; '
    'otherwise use the global default tariff.'
)
numbered(
    'For each damage, look for a matching Tariff Rule using the component code and repair code. '
    'If no exact match exists, try component only, or repair only.'
)
numbered(
    'Calculate labour hours, labour cost, and material cost from the matched rule or slab.'
)
numbered(
    'Apply the currency conversion factor if the estimate is in a currency other than USD '
    '(since all tariff rates are stored in USD).'
)
numbered(
    'Resolve the Charge Code and Tax Code automatically from the MnR Code Mapping table.'
)
numbered(
    'Store the calculated line item on the estimate, including the full breakdown '
    '(labour hours, labour rate, labour amount, material qty, material rate, material amount, ancillary).'
)

doc.add_paragraph()
para('The complete calculation for a single estimate line item is:', bold=True)

add_table(
    headers=['Field', 'Formula', 'Notes'],
    rows=[
        ['Labour Amount',    'Labour Hours × Labour Rate',                      'Both already in estimate currency'],
        ['Material Amount',  'Material Qty × Material Rate',                    'Both already in estimate currency'],
        ['Line Amount (net)','Qty × Unit Price',                                'Unit Price = Labour + Material + Ancillary'],
        ['Tax 1 Amount',     'Line Amount × Tax1 Rate ÷ 100',                  'e.g. SSCL — applies to net amount'],
        ['Tax 2 Amount',     '(Line Amount + Tax1) × Tax2 Rate ÷ 100',         'e.g. VAT — applies after Tax 1'],
        ['Gross Amount',     'Line Amount + Tax1 Amount + Tax2 Amount',         'Total per line including all taxes'],
    ],
    col_widths=[1.5, 2.5, 2.3],
    header_fill='1F4E79'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ESTIMATE TOTALS
# ══════════════════════════════════════════════════════════════════════════════
heading('5.  Estimate Totals', 1)

para(
    'After all line items are calculated, the system adds them up to produce the '
    'estimate header totals that are shown to the customer.',
    size=11
)

add_table(
    headers=['Estimate Total Field', 'How It Is Calculated'],
    rows=[
        ['Subtotal',          'Sum of all line amounts (net, before tax)'],
        ['SSCL Total',        'Sum of all Tax 1 amounts across all line items'],
        ['VAT Total',         'Sum of all Tax 2 amounts across all line items'],
        ['Total Tax',         'SSCL Total + VAT Total'],
        ['Grand Total',       'Subtotal + Total Tax'],
    ],
    col_widths=[2.0, 4.3],
    header_fill='2E75B6'
)

example_box(
    'Estimate with 3 Line Items (in LKR, exchange rate 1 USD = 300 LKR)',
    [
        '                     Net (LKR)   SSCL(2.5%)   VAT(18%)    Gross (LKR)',
        '  Line 1: Straighten panel  11,250.00     281.25     2,072.25    13,603.50',
        '  Line 2: Replace floor     31,200.00     780.00     5,724.00    37,704.00',
        '  Line 3: Weld crack         5,625.00     140.63     1,034.63     6,800.26',
        '  ─────────────────────────────────────────────────────────────────────',
        '  Subtotal                  48,075.00',
        '  SSCL Total                 1,201.88',
        '  VAT Total                  8,830.88',
        '  Grand Total               58,107.76  LKR',
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — CURRENCY CONVERSION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('6.  Multi-Currency Estimates', 1)

para(
    'All tariff rates are stored in USD. When you create an estimate for a '
    'local customer (e.g., a local agent or a domestic shipping line), '
    'you can choose to send the estimate in LKR (or any other currency). '
    'The system will convert all tariff rates using the exchange rate you specify.',
    size=11
)

heading('6.1  Choosing the Currency', 2)

para(
    'On the Create / Edit Estimate form, you select the Estimate Currency '
    '(USD, LKR, EUR, GBP, etc.). The system automatically looks up the '
    'exchange rate for the estimate date from the Exchange Rate master.',
    size=11
)

add_table(
    headers=['Scenario', 'Currency to Use', 'Exchange Rate'],
    rows=[
        ['Estimate for an overseas shipping line (Evergreen, MSC, Maersk)',
         'USD', '1.00 (no conversion needed)'],
        ['Estimate for a local agent or Sri Lankan company',
         'LKR', 'Auto-loaded from Exchange Rate master\n(e.g., 1 USD = 300.25 LKR)'],
        ['Estimate for a European entity',
         'EUR', 'Auto-loaded from Exchange Rate master\n(e.g., 1 USD = 0.91 EUR)'],
    ],
    col_widths=[2.5, 1.3, 2.5],
    header_fill='2E75B6'
)

heading('6.2  How Conversion Is Applied', 2)

para(
    'The conversion factor is applied when the system imports damages from the survey. '
    'All tariff amounts (labour rate, material rate, ancillary) are multiplied by the '
    'exchange rate so that every stored amount is already in the estimate\'s currency.',
    size=11
)

example_box(
    'Replacing a floor board — USD tariff converted to LKR',
    [
        'Tariff (in USD):',
        '  Labour Hours    : 2.00 hrs',
        '  Labour Rate     : USD 25.00 / hr',
        '  Material Qty    : 3 boards',
        '  Material Rate   : USD 18.00 / board',
        '',
        'Estimate Currency : LKR',
        'Exchange Rate     : 1 USD = 300.00 LKR',
        '',
        'After conversion (stored on estimate):',
        '  Labour Rate     : 25.00 × 300 = LKR 7,500.00 / hr',
        '  Labour Amount   : 2.00 × 7,500 = LKR 15,000.00',
        '  Material Rate   : 18.00 × 300 = LKR 5,400.00 / board',
        '  Material Amount : 3 × 5,400  = LKR 16,200.00',
        '  Line Total      : LKR 31,200.00',
    ]
)

box_note(
    'Once an estimate is sent to the customer, the exchange rate is locked and cannot be changed. '
    'This ensures the customer always sees the same amounts that were quoted.'
)

heading('6.3  What the Customer Sees', 2)
para(
    'When a customer opens the estimate in the Owner Portal or receives the estimate email, '
    'they see all amounts in the estimate\'s currency. If conversion was applied, '
    'a notice is shown, for example:',
    size=11
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run(
    '"All amounts are shown in LKR. '
    'Rates converted from USD at 1 USD = 300.0000 LKR as at 01 Jun 2025."'
)
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0, 112, 192)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — CHARGE CODES & TAXES
# ══════════════════════════════════════════════════════════════════════════════
heading('7.  Charge Codes and Tax Codes', 1)

para(
    'Each estimate line item must be linked to a Charge Code (for billing) '
    'and a Tax Code (for tax calculation). The system assigns these automatically '
    'using the MnR Code Charge Mapping table.',
    size=11
)

heading('7.1  MnR Code Charge Mapping', 2)

para(
    'The system finds the best matching Charge Code by looking up the '
    'Component Code and Repair Code combination. It uses a specificity rule — '
    'a more specific match wins over a general one:',
    size=11
)

add_table(
    headers=['Match Type', 'Specificity Score', 'Example'],
    rows=[
        ['Component Code + Repair Code matched', '3 (most specific)',  'PANEL + STR → "Hull Repair Charge"'],
        ['Component Code only matched',          '2',                  'PANEL + any → "Panel Work Charge"'],
        ['Repair Code only matched',             '1',                  'any + WLD → "Welding Charge"'],
        ['No code matched (wildcard)',           '0 (least specific)', 'any + any → "General MnR Charge"'],
    ],
    col_widths=[2.3, 1.5, 2.5],
    header_fill='2E75B6'
)

para(
    'Once the Charge Code is found, its linked Tax Code is used to '
    'determine the Tax 1 and Tax 2 rates for that line item.',
    size=11
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — CUSTOMER APPROVAL FLOW
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('8.  Customer Approval Flow', 1)

para(
    'After the estimate is created, it goes through an approval workflow before '
    'any repair work begins.',
    size=11
)

add_table(
    headers=['Status', 'Meaning', 'Who Sets It'],
    rows=[
        ['Draft',             'Estimate is being prepared. Not yet sent.', 'System (on creation)'],
        ['Sent',              'Estimate has been emailed to the customer.', 'Staff clicks "Send"'],
        ['Under Review',      'Customer has opened the portal and is reviewing.', 'System (on first portal access)'],
        ['Partially Approved','Customer has approved some lines but not all.', 'Customer (in portal)'],
        ['Approved',          'Customer has approved all line items. Repair can start.', 'Customer (in portal)'],
        ['Rejected',          'Customer has rejected the estimate.', 'Customer (in portal)'],
        ['Revised',           'A new version has been created. Original is superseded.', 'Staff'],
    ],
    col_widths=[1.5, 2.8, 2.0],
    header_fill='1F4E79'
)

heading('8.1  What the Customer Sees in the Portal', 2)

para(
    'The customer receives a unique link (no password needed) and sees a table with:',
    size=11
)
bullet('Component — which part of the container needs repair')
bullet('Repair Type — what kind of repair (e.g., Straighten, Replace, Weld, Paint)')
bullet('Labour — number of hours and the labour cost')
bullet('Materials — material cost for the repair (and ancillary charges if any)')
bullet('Line Total — total amount for that line')
bullet('Action — Accept or Reject button for each line item')

doc.add_paragraph()
para(
    'The customer can approve or reject each line individually. '
    'If all lines are approved, the estimate status becomes "Approved" and '
    'the repair work can begin.',
    size=11
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — FULL WORKED EXAMPLE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('9.  Full Worked Example', 1)

para(
    'Below is a complete end-to-end example showing how a 40ft container '
    'with three damages becomes a finished estimate.',
    size=11
)

heading('9.1  Container & Survey Details', 2)
add_table(
    headers=['Field', 'Value'],
    rows=[
        ['Container No.', 'TCKU4567890'],
        ['Size / Type',   '40ft General Purpose (GP)'],
        ['Customer',      'TransOcean Shipping Line'],
        ['Tariff Used',   'TransOcean Shipping — 2025 Tariff (customer-specific)'],
        ['Estimate Currency', 'USD (overseas shipping line)'],
        ['Exchange Rate', '1.00 (no conversion)'],
    ],
    col_widths=[2.0, 4.3],
    header_fill='375623'
)

heading('9.2  Damages Found in Survey', 2)
add_table(
    headers=['#', 'Location', 'Component', 'Damage', 'Repair', 'Qty', 'Dimensions'],
    rows=[
        ['1', 'Left Side', 'Panel',      'Dented',   'Straighten (STR)', '1',  'L: 18 in, W: 12 in'],
        ['2', 'Floor',     'Floor Board', 'Cracked', 'Replace (RPL)',    '3',   '—'],
        ['3', 'Door',      'Twist Lock',  'Broken',  'Replace (RPL)',    '8 (locks)', '—'],
    ],
    col_widths=[0.3, 0.8, 1.1, 0.9, 1.4, 0.8, 1.5],
    header_fill='2E75B6'
)

heading('9.3  Tariff Lookup Results', 2)
add_table(
    headers=['Damage', 'Tariff Rule / Slab Used', 'Labour Hrs', 'Labour Rate (USD)', 'Material Cost (USD)', 'Ancillary'],
    rows=[
        ['Panel Straighten',   'Rule: PANEL + STR',                    '1.50', '25.00', '0.00',  '0.00'],
        ['Floor Board Replace','Rule: FLOOR + RPL',                    '2.00', '25.00', '54.00', '0.00'],
        ['Twist Lock Replace', 'Item: LOCK-RPL, Slab (qty=8 → Slab2 + 3 extra)', '1.60', '25.00', '80.00', '0.00'],
    ],
    col_widths=[1.5, 2.3, 0.8, 1.1, 1.2, 0.8],
    header_fill='2E75B6'
)

heading('9.4  Calculated Estimate Line Items', 2)
add_table(
    headers=['#', 'Component / Repair', 'Labour Hrs', 'Labour Cost', 'Material Cost', 'Ancillary', 'Net (USD)', 'Tax (18%)', 'Gross (USD)'],
    rows=[
        ['1', 'Panel — Straighten',       '1.50', '37.50',  '0.00',  '0.00', '37.50', '6.75',  '44.25'],
        ['2', 'Floor Board — Replace',    '2.00', '50.00', '54.00',  '0.00', '104.00','18.72', '122.72'],
        ['3', 'Twist Lock — Replace (×8)','1.60', '40.00', '80.00',  '0.00', '120.00','21.60', '141.60'],
        ['',  'TOTALS',                   '5.10','127.50', '134.00', '0.00', '261.50','47.07', '308.57'],
    ],
    col_widths=[0.3, 1.7, 0.7, 0.8, 0.9, 0.7, 0.8, 0.7, 0.8],
    header_fill='1F4E79'
)

heading('9.5  Estimate Summary', 2)
add_table(
    headers=['Summary Line', 'Amount (USD)'],
    rows=[
        ['Subtotal (net)',          '261.50'],
        ['Tax (VAT 18%)',            '47.07'],
        ['Grand Total',             '308.57'],
        ['Total Labour Hours',        '5.10 hrs'],
        ['Total Labour Cost',       '127.50'],
        ['Total Material Cost',     '134.00'],
        ['Total Ancillary',           '0.00'],
    ],
    col_widths=[3.0, 1.5],
    header_fill='375623'
)

para(
    'This estimate is sent to TransOcean Shipping Line. The customer sees '
    'the breakdown of labour and materials for each repair, and can approve '
    'or reject individual lines through the portal.',
    size=11, italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — QUICK REFERENCE FORMULAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('10.  Quick Reference — All Formulas', 1)

add_table(
    headers=['Calculation', 'Formula'],
    rows=[
        ['Labour Amount (per line)',          'Labour Hours × Labour Rate'],
        ['Material Amount (per line)',         'Material Qty × Material Rate'],
        ['Unit Price',                         'Labour Amount + Material Amount + Ancillary'],
        ['Line Amount (net)',                   'Qty × Unit Price'],
        ['Tax 1 Amount (e.g. SSCL)',           'Line Amount × Tax1 Rate ÷ 100'],
        ['Tax 2 Amount (e.g. VAT)',            '(Line Amount + Tax1 Amount) × Tax2 Rate ÷ 100'],
        ['Gross Amount (per line)',             'Line Amount + Tax1 + Tax2'],
        ['Slab Extra Hours (if additional slab)',  '((Qty − Base Qty From) ÷ Add Qty) × Add Labour Hours'],
        ['Slab Extra Material',                '((Qty − Base Qty From) ÷ Add Qty) × Add Material Cost'],
        ['Currency Conversion (USD → other)',  'All USD amounts × Exchange Rate'],
        ['Estimate Subtotal',                  'SUM of all Line Amounts'],
        ['Estimate Grand Total',               'Subtotal + SUM of all Tax Amounts'],
    ],
    col_widths=[3.0, 3.3],
    header_fill='1F4E79'
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Container Yard Management System  ·  MnR Estimate Calculation Guide  ·  Internal Use Only'
)
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(150, 150, 150)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/user/container-yard/public/MnR_Estimate_Calculation_Guide.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
